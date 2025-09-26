import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from fpdf import FPDF
import os
from datetime import datetime

def generar_documentos(excel_path, carpeta_imagenes):
    carpeta_word = "documentos_word"
    carpeta_pdf = "documentos_pdf"
    archivo_reporte_general = "reporte_general.pdf"
    logo_path = "logo_sena.png"

    os.makedirs(carpeta_word, exist_ok=True)
    os.makedirs(carpeta_pdf, exist_ok=True)

    df = pd.read_excel(excel_path)

    # Validar evidencias
    faltantes = []
    for index, row in df.iterrows():
        evidencia = os.path.join(carpeta_imagenes, row["Evidencia"])
        if not os.path.exists(evidencia):
            faltantes.append(evidencia)
        else:
            df.at[index, "Evidencia"] = evidencia  # Actualiza ruta completa

    if faltantes:
        raise FileNotFoundError(f"Faltan evidencias: {faltantes}")

    # Generar documentos Word
    for index, row in df.iterrows():
        nombre = row["Nombre"]
        ficha = str(row["Ficha"])
        evidencia = row["Evidencia"]

        subcarpeta_word = os.path.join(carpeta_word, ficha)
        subcarpeta_pdf = os.path.join(carpeta_pdf, ficha)
        os.makedirs(subcarpeta_word, exist_ok=True)
        os.makedirs(subcarpeta_pdf, exist_ok=True)

        nombre_archivo_base = f"{ficha}_{nombre.replace(' ', '_')}"
        ruta_docx = os.path.join(subcarpeta_word, f"{nombre_archivo_base}.docx")

        doc = Document()
        doc.add_paragraph(f"FICHA: {ficha}")
        doc.add_paragraph(f"APRENDIZ:  {nombre}")
        doc.add_paragraph("EVIDENCIA CORREO")
        doc.add_picture(evidencia, width=Inches(5))
        doc.save(ruta_docx)

    # Convertir a PDF
    for ficha in df["Ficha"].astype(str).unique():
        subcarpeta_word = os.path.join(carpeta_word, ficha)
        subcarpeta_pdf = os.path.join(carpeta_pdf, ficha)
        convert(subcarpeta_word, subcarpeta_pdf)

    # Crear resumen por ficha
    agrupado = df.groupby("Ficha")
    for ficha, grupo in agrupado:
        ficha_str = str(ficha)
        subcarpeta_word = os.path.join(carpeta_word, ficha_str)
        resumen_path = os.path.join(subcarpeta_word, f"resumen_ficha_{ficha_str}.txt")
        with open(resumen_path, "w", encoding="utf-8") as resumen:
            resumen.write(f"Resumen de Ficha: {ficha_str}\n")
            resumen.write(f"Total de aprendices: {len(grupo)}\n\n")
            for nombre in grupo["Nombre"]:
                resumen.write(f"- {nombre}\n")

    # Crear reporte general PDF
    fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M")
    total_fichas = len(agrupado)
    total_aprendices = len(df)

    pdf = FPDF()
    pdf.add_page()

    if os.path.exists(logo_path):
        pdf.image(logo_path, x=10, y=8, w=30)
    pdf.set_xy(0, 35)

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Reporte de Aprendices enviados a cancelación por Formulario", ln=True, align="C")
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 10, f"Fecha de creación: {fecha_actual}", ln=True, align="C")
    pdf.ln(10)

    for ficha, grupo in agrupado:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, f"Ficha: {ficha}", ln=True)
        pdf.set_font("Arial", "", 11)
        for nombre in grupo["Nombre"]:
            pdf.cell(0, 8, f"- {nombre}", ln=True)
        pdf.ln(5)

    pdf.set_font("Arial", "B", 12)
    pdf.ln(10)
    pdf.cell(0, 10, "Resumen Final", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 8, f"Total de fichas procesadas: {total_fichas}", ln=True)
    pdf.cell(0, 8, f"Total de aprendices incluidos: {total_aprendices}", ln=True)

    pdf.output(archivo_reporte_general)

    return archivo_reporte_general
